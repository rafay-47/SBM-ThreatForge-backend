"""Word (.docx) export matching the web app structure."""


def export_word(model: dict, out_path: str) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    # ── Style overrides ───────────────────────────────────────────────────────
    _BLUE = RGBColor(0x2E, 0x74, 0xB5)
    h1_style = doc.styles["Heading 1"]
    h1_style.font.color.rgb = _BLUE
    h1_style.font.size = Pt(16)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _fmt_col(col: str) -> str:
        return " ".join(w.capitalize() for w in col.split("_"))

    def _shade_cell(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_table(heading: str, columns: list, rows: list, font_size: int = 9) -> None:
        if not rows:
            return
        doc.add_heading(heading, 1)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        # Header row
        for i, col in enumerate(columns):
            cell = table.rows[0].cells[i]
            cell.text = _fmt_col(col)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "428BCA")
        # Data rows
        for row_data in rows:
            row = table.add_row().cells
            for i, col in enumerate(columns):
                val = row_data.get(col, "") or ""
                if isinstance(val, list):
                    val = "\n".join(f"• {v}" for v in val)
                row[i].text = str(val)
                if row[i].paragraphs[0].runs:
                    row[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        doc.add_paragraph()  # spacing after table

    # ── Extract data ──────────────────────────────────────────────────────────
    title = model.get("title") or "Threat Model"
    description = model.get("description")
    assumptions = model.get("assumptions") or []
    assets = (model.get("assets") or {}).get("assets", [])
    arch = model.get("system_architecture") or {}
    data_flows = arch.get("data_flows", [])
    trust_bnd = arch.get("trust_boundaries", [])
    thr_src = arch.get("threat_sources", [])
    threats = (model.get("threat_list") or {}).get("threats", [])

    # ── Title ─────────────────────────────────────────────────────────────────
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = _BLUE
    doc.add_paragraph()

    # ── Architecture Diagram ──────────────────────────────────────────────────
    image_path = model.get("image_path")
    if image_path:
        from docx.shared import Cm
        from pathlib import Path

        p = Path(image_path)
        if p.exists():
            doc.add_heading("Architecture Diagram", 1)
            doc.add_picture(str(p), width=Cm(15))
            doc.add_paragraph()

    # ── Portrait sections ─────────────────────────────────────────────────────
    if description:
        doc.add_heading("Description", 1)
        doc.add_paragraph(description)
        doc.add_paragraph()

    if assumptions:
        _add_table(
            "Assumptions", ["assumption"], [{"assumption": a} for a in assumptions]
        )

    _add_table("Assets", ["type", "name", "description", "criticality"], assets)

    _add_table(
        "Data Flow", ["flow_description", "source_entity", "target_entity"], data_flows
    )

    _add_table(
        "Trust Boundary", ["purpose", "source_entity", "target_entity"], trust_bnd
    )

    _add_table("Threat Source", ["category", "description", "example"], thr_src)

    # ── Threat Catalog (landscape section) ────────────────────────────────────
    if threats:
        # Start a new landscape section
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )

        doc.add_heading("Threat Catalog", 1)
        cols = [
            "name",
            "stride_category",
            "description",
            "target",
            "impact",
            "likelihood",
            "mitigations",
        ]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for i, col in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = _fmt_col(col)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "428BCA")
        for t in threats:
            mitigations = t.get("mitigations") or []
            if isinstance(mitigations, list):
                mit_str = "\n".join(f"• {m}" for m in mitigations)
            else:
                mit_str = str(mitigations)
            vals = [
                t.get("name", ""),
                t.get("stride_category", ""),
                t.get("description", ""),
                t.get("target", ""),
                t.get("impact", ""),
                t.get("likelihood", ""),
                mit_str,
            ]
            row = table.add_row().cells
            for i, val in enumerate(vals):
                row[i].text = str(val) if val else ""
                if row[i].paragraphs[0].runs:
                    row[i].paragraphs[0].runs[0].font.size = Pt(8)

    doc.save(out_path)
